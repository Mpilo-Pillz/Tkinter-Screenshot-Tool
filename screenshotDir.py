from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from tkinter.filedialog import askopenfilename
from time import localtime, strftime
from docx import Document
from docx.shared import Inches
from os.path import expanduser
import time


import pyautogui
import os


class screenShotToolII:

    setNewImgDir = ''
    def __init__(self, master):
        frame = Frame(master)
        frame.pack()



        # pathString = askopenfilename(filetypes=[("Text files", "*.txt")])
        # if pathString != "":
        #     openFile = open(pathString, 'r')
        #     fileString = openFile.read()
        #     print(fileString)
        # root.destroy()

        # self.pathString = Button(frame, text="schoose")
        # self.pathString.pack(side=LEFT)

        self.quitButton = Button(frame, text="Stop", command=frame.quit)
        self.quitButton.pack(side=LEFT)

        self.printButton = Button(frame, text="Take Screen Shot", command=self.startScreenShotTool)
        self.printButton.pack(side=LEFT)

        self.saveDocButton = Button(frame, text="Save To Word Document", command=self.saveToDoc)
        self.saveDocButton.pack(side=LEFT)


        self.storePictureAtDir = Button(frame, text="Browse Image", command=self.saveImageToDir)
        self.storePictureAtDir.pack(side=RIGHT)

        self.imgDir = Text(frame, height=2, width=40)
        self.imgDir.pack(side=RIGHT)

        self.storeDocAtDir = Button(frame, text="Browse Document", command=self.saveDocToDir)
        self.storeDocAtDir.pack(side=RIGHT)

        self.docDir = Text(frame, height=2, width=40)
        self.docDir.pack(side=RIGHT)



    # print('Please click on the \"Take ScreenShot Button\"!')


    def startScreenShotTool(self):

        root.withdraw()
        time.sleep(1)
        todaysdate = str(strftime("%d %b %Y - %H %M %S"))

        # Save the screenshot in a folder and save the image using current date, time (hh:mm:ss)
        # img = pyautogui.screenshot('/home/mpilopillz/Pictures/MpiloToolPics/%s.png' % (todaysdate))
        img = pyautogui.screenshot(self.setNewImgDir + '/%s.png' % (todaysdate))
        root.update()
        root.deiconify()
        currentDir = self.setNewImgDir
        print(currentDir)
        print("mpilo")



    def saveToDoc(self):

        home = expanduser("~")
        # saveLocation = home + '/Documents/testscreen/'
        saveLocation = self.setNewDocDir + '/'
        # imgLocation= '%s/Pictures/MpiloToolPics/' % (home)
        imgLocation = self.setNewImgDir + '/'
        namedoc = str(strftime("%d %b %Y - %H %M"))
        doc = Document()
        # tables = doc.tables]




        for file in os.listdir(imgLocation):  # loop through all the files and folders for adding pictures
            if file.endswith(".png"):
                p = doc.add_paragraph()
                r = p.add_run()
                r.add_picture(imgLocation + file, width=Inches(6),height=Inches(4))

        doc.save('%s%s.docx' % (saveLocation, namedoc))


    def saveImageToDir(self):
        # global setNewImgDir
        selectImgFolder = filedialog.askdirectory()
        self.setNewImgDir = selectImgFolder
        print(selectImgFolder)
        self.imgDir.delete(1.0,END)
        self.imgDir.insert(INSERT, selectImgFolder)
        # return setNewImgDir

        # root.filename = filedialog.askopenfilename(initialdir=os.getcwd(), title="Select file",filetypes=[("Text Files", "*.txt")])

    def saveDocToDir(self):
        # global setNewImgDir
        selectDocFolder = filedialog.askdirectory()
        self.setNewDocDir = selectDocFolder
        print(selectDocFolder)
        self.docDir.delete(1.0,END)
        self.docDir.insert(INSERT, selectDocFolder)




        # current_img_Dir = ''
        #
        # def oneFunction(lists):
        #     global current_img_Dir
        #     word = random.choice(lists[category])
        #     current_word = word
        #
        # def anotherFunction():
        #     for letter in get_word():
        #         print("_", end=" ")
        #
        # def get_word():
        #     return current_word





root = Tk()
b = screenShotToolII(root)
root.title("Screen Shot Taker")
root.mainloop()