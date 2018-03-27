from tkinter import *
from time import localtime, strftime
from docx import Document
from docx.shared import Inches
from os.path import expanduser


import pyautogui
import os


class screenShotTool:

    def __init__(self, master):
        frame = Frame(master)
        frame.pack()

        self.quitButton = Button(frame, text="Stop", command=frame.quit)
        self.quitButton.pack(side=LEFT)

        self.printButton = Button(frame, text="Take Screen Shot", command=self.startScreenShotTool)
        self.printButton.pack(side=LEFT)

        self.saveDocButton = Button(frame, text="Save To Word Document", command=self.saveToDoc)
        self.saveDocButton.pack(side=LEFT)

    print('Please click on the \"Take ScreenShot Button\"!')
    def startScreenShotTool(self):
        todaysdate = str(strftime("%d %b %Y - %H %M %S"))

        # Save the screenshot in a folder and save the image using current date, time (hh:mm:ss)
        img = pyautogui.screenshot('/home/mpilopillz/Pictures/MpiloToolPics/%s.png' % (todaysdate))



    def saveToDoc(self):

        home = expanduser("~")
        saveLocation = home + '/Documents/testscreen/'
        imgLocation= '%s/Pictures/MpiloToolPics/' % (home)
        namedoc = str(strftime("%d %b %Y - %H %M"))
        doc = Document()
        # tables = doc.tables



        for file in os.listdir(imgLocation):  # loop through all the files and folders for adding pictures
            if file.endswith(".png"):
                p = doc.add_paragraph()
                r = p.add_run()
                r.add_picture(imgLocation + file, width=Inches(6),height=Inches(4))

        doc.save('%s%s.docx' % (saveLocation, namedoc))






root = Tk()
b = screenShotTool(root)
root.title("Screen Shot Taker")
root.mainloop()